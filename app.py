import streamlit as st
import pdfplumber
import pandas as pd
import os
import re
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

# Titre de l'application
st.title('Extraction automatisée de tâches pour la gestion des essais cliniques')
st.subheader('Datathon Défin n 12)

# Téléchargement du fichier PDF
uploaded_file = st.file_uploader("Télécharger un fichier PDF", type="pdf")

if uploaded_file is not None:
    # Lire le fichier PDF
    with pdfplumber.open(uploaded_file) as pdf:
        # Accéder à la page 11 (index 10, car l'indexation commence à 0)
        page = pdf.pages[10]
        
        # Extraire les tables de la page
        tables = page.extract_tables()
        
        # Vérifier si des tables ont été trouvées
        if tables:
            # Utiliser la première table trouvée
            table = tables[0]
            
            # Convertir les données du tableau en DataFrame Pandas
            df = pd.DataFrame(table)
            
            # Remplir les cellules vides avec la valeur précédente
            df.ffill(axis=1, inplace=True)
            
            # Afficher le DataFrame
            st.write("Tableau extrait de la page 11 du PDF :")
            st.dataframe(df)
            
            # Enregistrer le DataFrame en fichier CSV si nécessaire
            df.to_csv("output_table.csv", index=False)
            
            # Génération des documents Word
            st.write("Génération des documents Word...")

            # Fonction pour ajouter une case à cocher
            def add_checkbox(cell):
                checkbox = OxmlElement('w:sym')
                checkbox.set(qn('w:font'), 'Wingdings')
                checkbox.set(qn('w:char'), 'F0A8')  # Code Wingdings pour une case à cocher vide
                run = cell.paragraphs[0].add_run()
                run._r.append(checkbox)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Fonction pour ajouter une tâche dans le document
            def add_task(doc, task_description):
                table = doc.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                # Configurer la largeur des colonnes
                table.columns[0].width = Inches(0.3)
                table.columns[1].width = Inches(4.0)
                table.columns[2].width = Inches(2.0)

                # Ajouter la case à cocher
                add_checkbox(table.cell(0, 0))

                # Ajouter la description de la tâche
                task_cell = table.cell(0, 1)
                task_cell.text = task_description
                task_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                # Ajouter les champs pour l'heure et les initiales
                details_cell = table.cell(0, 2)
                details_cell.text = "Heure: __h__   Initiale: __/__"
                details_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Lire le fichier CSV dans un DataFrame
            df = pd.read_csv("output_table.csv")

            # Vérifier que le DataFrame contient au moins 3 lignes et 14 colonnes
            if df.shape[0] >= 3 and df.shape[1] >= 14:
                # Créer un CRF pour chaque jour pour les colonnes 3 à 14
                unique_key_counter = 0
                for column in range(3, 15):
                    cycle_info = str(df.iloc[0, column]).strip().replace('/', '-').replace('\\', '-')
                    day_info = str(df.iloc[1, column]).strip().replace('/', '-').replace('\\', '-')
                    
                    if cycle_info.lower() == 'nan' or day_info.lower() == 'nan':
                        st.write(f"Cycle info or day info is NaN at column {column}")
                        continue
                    
                    doc = Document()
                    doc.add_heading(f'Cycle: {cycle_info}', level=1)
                    doc.add_heading(f'Day: {day_info}', level=2)

                    for index, row in df.iterrows():
                        if index >= 3:
                            task_description = str(row[column])
                            is_subtitle = task_description.isupper() and not "X" in task_description
                            is_task = "X" in task_description

                            if is_subtitle:
                                # Nettoyer le sous-titre et l'ajouter au document
                                cleaned_description = re.sub(r'\(\d+\)', '', task_description).strip()
                                doc.add_heading(cleaned_description, level=2)
                            elif is_task:
                                # Ajouter la tâche au document
                                add_task(doc, row[0])

                    # Définir le chemin complet pour enregistrer le fichier
                    filename = f'CRF_{cycle_info.replace(" ", "_")}_{day_info.replace(" ", "_")}.docx'
                    dossier_sortie = "generated_docs"  # Dossier pour enregistrer les fichiers générés
                    chemin_complet = os.path.join(dossier_sortie, filename)

                    # Vérifier et créer le dossier de sortie si nécessaire
                    os.makedirs(dossier_sortie, exist_ok=True)

                    # S'assurer que le nom du fichier n'est pas trop long
                    if len(chemin_complet) > 255:
                        chemin_complet = chemin_complet[:255]

                    try:
                        doc.save(chemin_complet)
                        st.write(f"Document généré: {filename}")
                        
                        # Préparer le document pour le téléchargement
                        with open(chemin_complet, "rb") as f:
                            st.download_button(
                                label=f"Télécharger {filename}",
                                data=f.read(),
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_button_{unique_key_counter}"
                            )
                            unique_key_counter += 1
                    except Exception as e:
                        st.write(f"Impossible de sauvegarder le document {filename}: {e}")
            else:
                st.write("Le DataFrame ne contient pas assez de lignes ou de colonnes pour procéder.")
        else:
            st.write("Aucune table trouvée à la page 11")
